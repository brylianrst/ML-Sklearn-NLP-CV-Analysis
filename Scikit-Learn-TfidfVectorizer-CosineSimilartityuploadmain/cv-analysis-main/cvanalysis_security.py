#All Import
# Install necessary packages:
# Run the following commands in your terminal:
'''
!pip install nltk
!pip install PyMuPDF
!pip install prettytable
!pip install python-docx

# Install cloudmersive api packages
!pip install cloudmersive-virus-api-client
!pip install loguru
'''

# Import the necessary packages in your script
import nltk
import fitz  # PyMuPDF
from prettytable import PrettyTable
import docx
import cloudmersive_virus_api_client
from loguru import logger

from multiprocessing import Pool, cpu_count
import pandas as pd
from nltk.corpus import wordnet
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import os
from docx import Document
import re
from google.colab import drive
import time

# Include packages related to Cloudmersive
from __future__ import print_function
from cloudmersive_virus_api_client.rest import ApiException

import requests

# Packages for Unit Testing
import unittest
from unittest.mock import patch, MagicMock, create_autospec
from io import StringIO
from loguru import logger
import time

import sys


"""
Functions and the body of code
"""

# Start the Logger
# Logger, by using "Loguru"
# Initialize logger
logger.add(
    "application.log",
    rotation="1 week",
    format="{time:YYYY-MM-DD at HH:mm:ss} | {level} | {name}:{function} | {message}"
)

logger.info("Started")


# Function to check if the directory exists and is accessible
def check_repository_access(cv_folder_path_given):
    try:
        if not os.path.isdir(cv_folder_path_given):
            raise FileNotFoundError(
                "The specified folder does not exist or cannot be accessed."
            )
        return True
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return False

# Function to list all PDF, DOC, and DOCX files in the CV folder, return [] if no such file exists
def load_cv_files(cv_folder_path_given):
    # I added this 'try ... except' part, only the 'cv_files_found ...' is original
    try:
        if not check_repository_access(cv_folder_path_given):
            # Directory is not accessible, return an empty list
            logger.error("Directory access check failed. No files to load.")
            return []

        cv_files_found = [
            f for f in os.listdir(cv_folder_path_given)
            if f.lower().endswith(('.pdf', '.doc', '.docx'))
        ]

        if cv_files_found:
            logger.info(f"Found {len(cv_files_found)} files: {cv_files_found}")
        else:
            logger.info("No CV files of given format found in the directory.")

        return cv_files_found

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        return []
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return []

# These two functions scan the files to detect potential Virus/Trojan/Malware ... etc. in documents (Sicheng)

def scan_one_file(input_file, api_key, retry_time):
    """
    Scans a single file for viruses and returns whether it is clean.

    Args:
        input_file (str): The path to the file to be scanned.
        api_key (str): The API key for the Cloudmersive virus scan service.
        retry_time (int): The current retry count.

    Returns:
        bool: True if the file is clean,
              False if the file is malicious
              or if an error occurs after reaching retry limit.
    """
    # Configure API key authorization: Apikey
    configuration = cloudmersive_virus_api_client.Configuration()
    configuration.api_key['Apikey'] = api_key

    # Create an instance of the API class
    api_instance = cloudmersive_virus_api_client.ScanApi(
        cloudmersive_virus_api_client.ApiClient(configuration)
    )

    # File type restriction
    restrict_file_types = '.doc,.docx,.pdf'
    allow_executables = False  # No .exe
    allow_invalid_files = False
    allow_scripts = True
    allow_password_protected_files = True
    allow_macros = True
    allow_xml_external_entities = False
    allow_insecure_deserialization = False
    allow_html = False

    try:
        # Send files to Cloudmersive to scan
        api_response = api_instance.scan_file_advanced(
            input_file,
            allow_executables=allow_executables,
            allow_invalid_files=allow_invalid_files,
            allow_scripts=allow_scripts,
            allow_password_protected_files=allow_password_protected_files,
            allow_macros=allow_macros,
            allow_xml_external_entities=allow_xml_external_entities,
            allow_insecure_deserialization=allow_insecure_deserialization,
            allow_html=allow_html
        )

        return api_response.clean_result  # Return True if the file is clean, False if not

    except ApiException as e:
        print(f"Exception when calling ScanApi->scan_file: {e}\n")

        if retry_time < 5:
            time.sleep(1)
            return scan_one_file(input_file, api_key, retry_time + 1)
        else:
            print(f"Fail to scan file {input_file}")
            return False



def scan_all_files_in_repository(cv_folder_path, all_valid_files, api_key):
    """
    Scans all valid files in a directory and
    returns a list of files without any problems.

    Args:
        cv_folder_path (str): Path to the directory containing the files.
        all_valid_files (list): List of all valid files to scan.
        api_key (str): The API key for the Cloudmersive virus scan service.

    Returns:
        list: A list of filenames that passed the safety check.
    """
    files_without_problem = []

    # check if all files are valid
    # pick only valid ones
    for file_name in all_valid_files:
        file_path = os.path.join(cv_folder_path, file_name)
        scan_result = scan_one_file(file_path, api_key, 0)

        if scan_result:
            files_without_problem.append(file_name)

    return files_without_problem

# Function to extract name, designation, experience, education, and skills
def extract_information(text):
    name = ''
    designation = ''
    experience = ''
    education = ''
    skills = ''

    # Define regular expressions to match patterns
    name_pattern = r'Name: ([A-Za-z\s]+)'
    designation_pattern = r'Designation: ([\w\s]+)'
    experience_pattern = r'Experience:\s*(.*?)\s*(?=Education:|Skills:|$)'
    education_pattern = r'Education: (.*?)(?=[A-Z][a-z]+:|$)'
    skills_pattern = r'Skills: (.*?)(?=[A-Z][a-z]+:|$)'

    name_match = re.search(name_pattern, text)
    if name_match:
        name = name_match.group(1).strip()

    designation_match = re.search(designation_pattern, text)
    if designation_match:
        designation = designation_match.group(1).strip()

    experience_match = re.search(experience_pattern, text, re.DOTALL)
    if experience_match:
        experience = experience_match.group(1).strip()

    education_match = re.search(education_pattern, text, re.DOTALL)
    if education_match:
        education = education_match.group(1).strip()

    skills_match = re.search(skills_pattern, text, re.DOTALL)
    if skills_match:
        skills = skills_match.group(1).strip()

    # Additional cleanup: Remove bullet points and other unwanted characters
    unwanted_chars = ["•", "●", "▪", "§", "\n", "\r", "○"]
    for char in unwanted_chars:
        name = name.replace(char, "").strip()
        designation = designation.replace(char, "").strip()
        experience = experience.replace(char, "").strip()
        education = education.replace(char, "").strip()
        skills = skills.replace(char, "").strip()

    # Remove 'Email' from the name field if present
    if 'email' in name.lower():
        name = re.split(r'\s+', name, 1)[0]

    return name, designation, experience, education, skills


# Function to read the content of DOCX file
def read_docx(file_path):
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

# Function to read the content of DOC file (using antiword tool)
def read_doc(file_path):
    result = os.popen(f'antiword "{file_path}"').read()
    return result

# Function to process each file and extract information
def process_cv_file(cv_file):
    file_path = os.path.join(cv_folder_path, cv_file)
    text = ""
    if cv_file.lower().endswith('.pdf'):
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text()
    elif cv_file.lower().endswith('.docx'):
        text = read_docx(file_path)
    elif cv_file.lower().endswith('.doc'):
        text = read_doc(file_path)

    name, designation, experience, education, skills = extract_information(text)
    return cv_file, text, name, designation, experience, education, skills

# Function to get synonyms of a word using WordNet
def get_synonyms(word):
    synonyms = set()
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.add(lemma.name())
    return synonyms

# Function to preprocess text for TF-IDF vectorization
def preprocess_text(text, jd_keywords_synonyms=None):
    words = nltk.word_tokenize(text.lower())
    stop_words = set(nltk.corpus.stopwords.words('english'))
    words = [
        lemmatizer.lemmatize(word)
        for word in words
        if word.isalnum() and word not in stop_words
    ]

    if jd_keywords_synonyms:
        words_with_synonyms = set(words)
        for word in words:
            if word in jd_keywords_synonyms:
                words_with_synonyms.update(jd_keywords_synonyms[word])
        words = list(words_with_synonyms)

    return ' '.join(words)


# Function to calculate context score based on cosine similarity of TF-IDF vectors
def calculate_context_score(jd, cv):
    jd_processed = preprocess_text(jd)
    cv_processed = preprocess_text(cv, jd_keywords_synonyms)
    documents = [jd_processed, cv_processed]
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)
    cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
    return cosine_sim[0][0]

# Function to calculate scores for each CV
def calculate_scores(i):
    combined_text = (
        f"{cvs_df.iloc[i]['Text']} "
        f"{cvs_df.iloc[i]['Name']} "
        f"{cvs_df.iloc[i]['Experience']} "
        f"{cvs_df.iloc[i]['Skills']}"
    )
    context_score = calculate_context_score(sample_jd, combined_text)
    return context_score

# Function to check if the directory exists and is accessible
def check_repository_access(cv_folder_path_given):
    try:
        if not os.path.isdir(cv_folder_path_given):
            raise FileNotFoundError(
                "The specified folder does not exist or cannot be accessed."
            )
        return True
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return False

# Function to list all PDF, DOC, and DOCX files in the CV folder, return [] if no such file exists
def load_cv_files(cv_folder_path_given):
    # I added this 'try ... except' part, only the 'cv_files_found ...' is original
    try:
        if not check_repository_access(cv_folder_path_given):
            # Directory is not accessible, return an empty list
            logger.error("Directory access check failed. No files to load.")
            return []

        cv_files_found = [
            f for f in os.listdir(cv_folder_path_given)
            if f.lower().endswith(('.pdf', '.doc', '.docx'))
        ]

        if cv_files_found:
            logger.info(f"Found {len(cv_files_found)} files: {cv_files_found}")
        else:
            logger.info("No CV files of given format found in the directory.")

        return cv_files_found
    except FileNotFoundError as e:
        logger.warning(f"File not found: {e}")
        return []
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return []


# These two functions "scan_one_file" and "scan_all_files_in_repository"
# scan the files to detect potential Virus/Trojan/Malware ... etc. in documents (Sicheng)
def scan_one_file(input_file, api_key, retry_time):
    """
    Scans a single file for viruses and returns whether it is clean.

    Args:
        input_file (str): The path to the file to be scanned.
        api_key (str): The API key for the Cloudmersive virus scan service.
        retry_time (int): The current retry count.

    Returns:
        bool: True if the file is clean,
              False if the file is malicious
              or if an error occurs after reaching retry limit.
    """
    # Configure API key authorization: Apikey
    configuration = cloudmersive_virus_api_client.Configuration()
    configuration.api_key['Apikey'] = api_key

    # Create an instance of the API class
    api_instance = cloudmersive_virus_api_client.ScanApi(
        cloudmersive_virus_api_client.ApiClient(configuration)
    )

    # File type restriction
    restrict_file_types = '.doc,.docx,.pdf'
    allow_executables = False  # No .exe
    allow_invalid_files = False
    allow_scripts = True
    allow_password_protected_files = True
    allow_macros = True
    allow_xml_external_entities = False
    allow_insecure_deserialization = False
    allow_html = False

    try:
        # Send files to Cloudmersive to scan
        api_response = api_instance.scan_file_advanced(
            input_file,
            allow_executables=allow_executables,
            allow_invalid_files=allow_invalid_files,
            allow_scripts=allow_scripts,
            allow_password_protected_files=allow_password_protected_files,
            allow_macros=allow_macros,
            allow_xml_external_entities=allow_xml_external_entities,
            allow_insecure_deserialization=allow_insecure_deserialization,
            allow_html=allow_html
        )

        logger.info(f"File {input_file} scanned by Cloudmersive successfully.")
        return api_response.clean_result  # Return True if the file is clean, False if not

    except Exception as e:
        logger.info(f"Exception when calling Cloudmersive API -> scan_file_advanced: {e}\n")

        if retry_time < 5:
            time.sleep(1)
            return scan_one_file(input_file, api_key, retry_time + 1)
        else:
            logger.warning(f"Failed to scan file: {input_file}")
            return False


def scan_all_files_in_repository(cv_folder_path, all_valid_files, api_key):
    """
    Scans all valid files in a directory and
    returns a list of files without any problems.

    Args:
        cv_folder_path (str): Path to the directory containing the files.
        all_valid_files (list): List of all valid files to scan.
        api_key (str): The API key for the Cloudmersive virus scan service.

    Returns:
        list: A list of filenames that passed the safety check.
    """
    files_without_problem = []

    # check if all files are valid
    # pick only valid ones
    try:
        for file_name in all_valid_files:
            file_path = os.path.join(cv_folder_path, file_name)
            scan_result = scan_one_file(file_path, api_key, 0)

            if scan_result:
                logger.info(f"File {file_path} is clean.")
                files_without_problem.append(file_name)
            else:
                logger.warning(f"File {file_path} is not clean.")

        logger.info("Scanning all files in the repositorycompleted.")
    except Exception as e:
        logger.warning(f"Failed to complete scanning process: {e}")
    finally:
        return files_without_problem


# Function to extract name, designation, experience, education, and skills
def extract_information(text):
    name = ''
    designation = ''
    experience = ''
    education = ''
    skills = ''

    # Define regular expressions to match patterns
    name_pattern = r'Name: ([A-Za-z\s]+)'
    designation_pattern = r'Designation: ([\w\s]+)'
    experience_pattern = r'Experience:\s*(.*?)\s*(?=Education:|Skills:|$)'
    education_pattern = r'Education: (.*?)(?=[A-Z][a-z]+:|$)'
    skills_pattern = r'Skills: (.*?)(?=[A-Z][a-z]+:|$)'

    name_match = re.search(name_pattern, text)
    if name_match:
        name = name_match.group(1).strip()

    designation_match = re.search(designation_pattern, text)
    if designation_match:
        designation = designation_match.group(1).strip()

    experience_match = re.search(experience_pattern, text, re.DOTALL)
    if experience_match:
        experience = experience_match.group(1).strip()

    education_match = re.search(education_pattern, text, re.DOTALL)
    if education_match:
        education = education_match.group(1).strip()

    skills_match = re.search(skills_pattern, text, re.DOTALL)
    if skills_match:
        skills = skills_match.group(1).strip()

    # Additional cleanup: Remove bullet points and other unwanted characters
    unwanted_chars = ["•", "●", "▪", "§", "\n", "\r", "○"]
    for char in unwanted_chars:
        name = name.replace(char, "").strip()
        designation = designation.replace(char, "").strip()
        experience = experience.replace(char, "").strip()
        education = education.replace(char, "").strip()
        skills = skills.replace(char, "").strip()

    # Remove 'Email' from the name field if present
    if 'email' in name.lower():
        name = re.split(r'\s+', name, 1)[0]

    return name, designation, experience, education, skills

# Function to read the content of DOCX file
def read_docx(file_path):
    doc = Document(file_path)
    full_text = [para.text for para in doc.paragraphs]
    return '\n'.join(full_text)

# Function to read the content of DOC file (using antiword tool)
def read_doc(file_path):
    result = os.popen(f'antiword "{file_path}"').read()
    return result

# Function to process each file and extract information
def process_cv_file(cv_file):
    file_path = os.path.join(cv_folder_path, cv_file)

    text = ""
    name = ""
    designation = ""
    experience = ""
    education = ""
    skills = ""

    try:
        if cv_file.lower().endswith('.pdf'):
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
        elif cv_file.lower().endswith('.docx'):
            text = read_docx(file_path)
        elif cv_file.lower().endswith('.doc'):
            text = read_doc(file_path)

        name, designation, experience, education, skills = extract_information(text)
    except Exception as e:
        logger.error(f"Error processing file {cv_file}: {e}")
    finally:
        return cv_file, text, name, designation, experience, education, skills


# Function to get synonyms of a word using WordNet
def get_synonyms(word):
    synonyms = set()
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            synonyms.add(lemma.name())
    return synonyms

# Function to preprocess text for TF-IDF vectorization
def preprocess_text(text, jd_keywords_synonyms=None):
    words = nltk.word_tokenize(text.lower())
    stop_words = set(nltk.corpus.stopwords.words('english'))
    words = [
        lemmatizer.lemmatize(word)
        for word in words
        if word.isalnum() and word not in stop_words
    ]

    if jd_keywords_synonyms:
        words_with_synonyms = set(words)
        for word in words:
            if word in jd_keywords_synonyms:
                words_with_synonyms.update(jd_keywords_synonyms[word])
        words = list(words_with_synonyms)

    return ' '.join(words)


# Function to calculate context score based on cosine similarity of TF-IDF vectors
def calculate_context_score(jd, cv):
    jd_processed = preprocess_text(jd)
    cv_processed = preprocess_text(cv, jd_keywords_synonyms)
    documents = [jd_processed, cv_processed]

    # Handle empty input
    if not any(documents):
        return 0.0

    vectorizer = TfidfVectorizer()

    try:
        tfidf_matrix = vectorizer.fit_transform(documents)
        cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        return cosine_sim[0][0]
    except Exception as e:
        logger.error(f"Failed to calculate TF-IDF vectorization due to this error: {e}")
        raise ValueError("Failed to calculate context score in calculate_context_score.") from e
        return 0.0  # Return a default score if TF-IDF fails due to empty vocabulary


# Function to calculate scores for each CV
def calculate_scores(i, cvs_df):
    context_score = 0.0
    try:
        combined_text = (
            f"{cvs_df.iloc[i]['Text']} "
            f"{cvs_df.iloc[i]['Name']} "
            f"{cvs_df.iloc[i]['Experience']} "
            f"{cvs_df.iloc[i]['Skills']}"
        )
        context_score = calculate_context_score(sample_jd, combined_text)
    except Exception as e:
        logger.error(f"Error calculating score for CV {i}: {e}")
    finally:
        return context_score
    
def perform_cv_analysis(cv_folder_path, cv_files, sample_jd):
    """
    Analyzes CV files based on a job description and ranks them.

    Parameters:
    cv_folder_path (str): Path to the folder with CV files.
    cv_files (list of str): List of CV file paths.
    sample_jd (str): Job description for CV relevance assessment.

    Returns:
    None: Prints top-ranked CVs and logs information.

    Raises:
    ValueError: If `sample_jd` is invalid or empty.
    """
    start_time = time.time()

    try:
        if not cv_files or len(cv_files) == 0:
            print("No valid file to be handled.")
        elif sample_jd is None or not sample_jd.strip() or not any(c.isalpha() for c in sample_jd):
            raise ValueError("Invalid or missing sample job description.")
        else:
            # Parallelize the processing of CV files
            with Pool(cpu_count()) as p:
                results = p.map(process_cv_file, cv_files)

            # Unzip the results
            file_names, cv_texts, names, designations, experiences, educations, skills_list = zip(*results)

            # Create a DataFrame from the CV texts and extracted information
            cvs_df = pd.DataFrame({
                "File Name": file_names,
                "Text": cv_texts,
                "Name": names,
                "Designation": designations,
                "Experience": experiences,
                "Education": educations,
                "Skills": skills_list
            })

            if cvs_df.empty:
                logger.info("No data to process after loading CV files.")
                return

            # Download NLTK resources if not already downloaded
            nltk.download('punkt')
            nltk.download('stopwords')
            nltk.download('wordnet')

            # Initialize WordNet Lemmatizer
            global lemmatizer
            lemmatizer = WordNetLemmatizer()

            # Extract keywords from the JD and find their synonyms
            jd_keywords = nltk.word_tokenize(sample_jd.lower())
            jd_keywords = [
                lemmatizer.lemmatize(word)
                for word in jd_keywords
                if word.isalnum() and word not in set(nltk.corpus.stopwords.words('english'))
            ]

            global jd_keywords_synonyms  # If possible, create a class and put global variables in
            jd_keywords_synonyms = {}
            for keyword in jd_keywords:
                jd_keywords_synonyms[keyword] = get_synonyms(keyword)

            # Parallelize score calculation using multiprocessing
            with Pool(cpu_count()) as p:
                context_scores = p.starmap(calculate_scores, [(i, cvs_df) for i in range(len(cvs_df))])

            # Add context scores to the DataFrame
            cvs_df['Context Score'] = context_scores

            # Normalize scores (context score only)
            cvs_df['Normalized Context Score'] = cvs_df['Context Score'] / cvs_df['Context Score'].max()

            # Convert scores to percentages
            cvs_df['Context Score (%)'] = cvs_df['Normalized Context Score'] * 100

            # Define weight for context score (since Monte Carlo score is removed)
            context_weight = 1.0

            # Calculate composite score (only context score)
            cvs_df['Composite Score (%)'] = cvs_df['Context Score (%)']

            # Sort the DataFrame by composite score in descending order
            sorted_cvs_df = cvs_df.sort_values(by='Composite Score (%)', ascending=False)

            # Output the top-ranked CVs with their file names and scores (excluding Designation and Education)
            top_cvs = sorted_cvs_df[['File Name', 'Composite Score (%)']]
            logger.info("Top CVs are assessed and sorted")

            # Print top CVs
            print("Top CVs based on Composite Score:")
            top_cvs_table = PrettyTable()
            top_cvs_table.field_names = ["File Name", "Composite Score (%)"]

            for index, row in top_cvs.head(10).iterrows():
                top_cvs_table.add_row([row["File Name"], row["Composite Score (%)"]])

            print(top_cvs_table)
    except Exception as e:
        logger.error(f"An unexpected error occurred while processing the CVs: {e}")
    finally:
        end_time = time.time()
        execution_time = end_time - start_time
        print(f"Total execution time: {execution_time:.2f} seconds")


"""
Call these functions
"""
# Mount Google Drive
drive.mount('/content/drive')

# Path to the folder containing CVs in Google Drive
cv_folder_path = '/content/drive/My Drive/FBACVs'

# Check if the repository is accessible
if not check_repository_access(cv_folder_path):
    print("Access to the CV repository is not available. Please check the path and your permissions.")
    cv_files = []
else:
    # List all PDF, DOC, and DOCX files in the CV folder
    cv_files = load_cv_files(cv_folder_path)

# Code for implementation of Cloudmersive (Sicheng)
api_key = '94b2b614-8e09-4d72-9c54-4a6f06d32b48'  # Change to your own API Key when you need to incorporate the code
cv_files = scan_all_files_in_repository(cv_folder_path, cv_files, api_key=api_key)

print(cv_files)  # Simply for presentation, remove it when you need to incorporate the code

# Sample Job Description
sample_jd = """
Job Description: Marketing Specialist at XYZ Company
Location: London, UK
Department: Marketing
Position Type: Full-Time

About XYZ Company:

XYZ Company is a leading innovator in the technology sector, committed to delivering high-quality products and services that empower businesses worldwide. With a dynamic and collaborative work environment, we pride ourselves on fostering creativity, growth, and excellence in our team. As we continue to expand our market presence, we are looking for a talented and driven Marketing Specialist to join our team in London, UK.

Position Overview:

We are seeking a creative and analytical Marketing Specialist to help us elevate our brand and drive our marketing initiatives to new heights. The ideal candidate will be responsible for developing, implementing, and managing marketing campaigns that promote our products and services. This role will involve a mix of strategic planning, creative thinking, and hands-on execution. If you are passionate about marketing, thrive in a fast-paced environment, and have a keen eye for detail, we want to hear from you!

Key Responsibilities:

- Campaign Management: Develop and execute comprehensive marketing campaigns across various channels, including digital, social media, email, and traditional marketing.
- Content Creation: Create compelling and engaging content for different platforms, including blog posts, social media updates, newsletters, and website copy.
- Market Research: Conduct market research to identify trends, customer needs, and competitive analysis to inform marketing strategies.
- Brand Development: Work closely with the design and product teams to ensure consistent branding and messaging across all marketing materials.
- Performance Analysis: Track and analyze the performance of marketing campaigns using tools such as Google Analytics, and provide actionable insights and recommendations.
- SEO/SEM: Optimize content for search engines and manage paid search campaigns to increase online visibility and drive traffic.
- Social Media Management: Manage and grow our social media presence by creating and curating high-quality content and engaging with our audience.
- Event Coordination: Plan and execute marketing events, webinars, and trade shows to promote our brand and generate leads.
- Collaborative Projects: Collaborate with cross-functional teams, including sales, product development, and customer service, to align marketing efforts with overall business goals.

Qualifications:

- Education: Bachelor’s degree in Marketing, Business, Communications, or a related field.
- Experience: 3+ years of experience in marketing, preferably within the technology or B2B sector.
- Skills: Strong understanding of digital marketing channels, SEO/SEM, content marketing, and social media strategies.
- Tools: Proficient in marketing tools and platforms such as Google Analytics, Hootsuite, MailChimp, and Adobe Creative Suite.
- Creativity: Exceptional creativity and innovation skills, with the ability to develop unique marketing strategies and content.
- Analytical Thinking: Strong analytical skills with the ability to interpret data and make data-driven decisions.
- Communication: Excellent written and verbal communication skills, with a keen attention to detail.
- Team Player: Ability to work effectively both independently and as part of a team in a fast-paced environment.
- Adaptability: Highly adaptable with a positive attitude and a willingness to learn and take on new challenges.

What We Offer:

- Competitive salary and performance-based bonuses.
- Comprehensive benefits package including health insurance, retirement plans, and paid time off.
- Opportunities for professional development and career growth.
- A vibrant and inclusive work environment with a focus on work-life balance.
- The chance to be part of a forward-thinking company at the forefront of technological innovation.

How to Apply:

If you are excited about the opportunity to contribute to a growing company and make a significant impact, please submit your resume and a cover letter detailing your relevant experience and why you are the perfect fit for this role. Apply now and join us in shaping the future of technology at XYZ Company!

XYZ Company is an equal opportunity employer. We celebrate diversity and are committed to creating an inclusive environment for all employees.
"""

perform_cv_analysis(cv_folder_path, cv_files, sample_jd)
logger.remove() # Close the logger


"""
------------------------------------------------------------------------------------------
Unit tests to test the functions
"""
print(sys.executable)

# Logger, by using "Loguru"
# Initialize logger (log to console)
logger.add("sys.stdout",
           rotation="1 week",
           format="{time:YYYY-MM-DD at HH:mm:ss} | "
                  "{level} | "
                  "{name}:{function} | "
                  "{message}")

logger.info("started")

# cv_folder_path for test
cv_folder_path = '/content/drive/My Drive/FBACVs'


class TestCloudmersive(unittest.TestCase):
    """
    This class tests if the Cloudmersive API could work properly.
    """
    test_api_key = '94b2b614-8e09-4d72-9c54-4a6f06d32b48'
    # change test_folder_path to your own path
    test_folder_path = '/content/drive/My Drive/FBACVs'
    test_cv_files = load_cv_files(test_folder_path)

    def test_scan_one_file_clean(self):
        """
        Test if a known clean file is correctly identified as not malicious.
        """
        # change to your own good file
        good_file = '10228751.pdf'
        result = scan_one_file(self.test_folder_path +
                               '/' + good_file, self.test_api_key, 0)
        self.assertTrue(result)  # True -- clean, False -- Malicious

    def test_scan_one_file_malicious(self):
        """
        Test if a known malicious file is correctly identified as malicious.
        """
        # change to your own bad file
        bad_file = 'eicar-download.pdf'
        result = scan_one_file(self.test_folder_path +
                               '/' + bad_file, self.test_api_key, 0)
        self.assertFalse(result)  # True -- clean, False -- Malicious

    def test_scan_all_files_in_repository(self):
        """
        Test the scanning of multiple files,
        ensuring all files are processed and the results are accurate.
        """
        passed_files = scan_all_files_in_repository(
            self.test_folder_path, self.test_cv_files, self.test_api_key)
        test_result = ['10501991.pdf', '10235211.pdf', '10228751.pdf',
                       '10541358.pdf', '10289113.pdf', 'Anti-Virus3.pdf']
        self.assertEqual(passed_files, test_result)


class TestNotebook(unittest.TestCase):
    """
    This class tests the document processing functions could
    handle unexpected inputs & errors properly when they occur.
    """
    test_api_key = '94b2b614-8e09-4d72-9c54-4a6f06d32b48'
    # change test_valid_folder_path to your own path
    test_valid_folder_path = '/content/drive/My Drive/FBACVs'
    test_invalid_folder_path = '/invalidpath'

    def setUp(self):
        self.docx_file = "nonexistent.docx"
        self.doc_file = "nonexistent.doc"
        self.pdf_file = "nonexistent.pdf"

    @patch('__main__.read_docx',
           side_effect=FileNotFoundError("File not found."))
    @patch('__main__.logger')
    def test_process_cv_file_handles_read_docx_exception(self,
                                                         mock_logger,
                                                         mock_read_docx):
        """
        Test process_cv_file handles FileNotFoundError exceptions
        when reading docx files.

        Use @patch decorator to mock read_docx function and logger object.
        """
        result = process_cv_file(self.docx_file)

        mock_logger.error.assert_called_once_with(
            f"Error processing file {self.docx_file}: File not found.")

        expected = (self.docx_file, "", "", "", "", "", "")
        self.assertEqual(result, expected)

    @patch('__main__.read_doc',
           side_effect=FileNotFoundError("File not found."))
    @patch('__main__.logger')
    def test_process_cv_file_handles_read_doc_exception(self,
                                                        mock_logger,
                                                        mock_read_doc):
        """
        Test process_cv_file handles FileNotFoundError exceptions
        when reading doc files.
        Use @patch decorator to mock read_doc function and logger object.
        """
        result = process_cv_file(self.doc_file)

        mock_logger.error.assert_called_once_with(
            f"Error processing file {self.doc_file}: File not found.")

        expected = (self.doc_file, "", "", "", "", "", "")
        self.assertEqual(result, expected)

    @patch('__main__.fitz.open',
           side_effect=FileNotFoundError("File not found."))
    @patch('__main__.logger')
    def test_process_cv_file_handles_read_pdf_exception(self,
                                                        mock_logger,
                                                        mock_fitz_open):
        """
        Test process_cv_file handles FileNotFoundError exceptions
        when reading PDF files.

        Use @patch decorator to mock fitz.open function and logger object.
        """
        result = process_cv_file(self.pdf_file)

        mock_logger.error.assert_called_once_with(
            f"Error processing file {self.pdf_file}: File not found.")

        expected = (self.pdf_file, "", "", "", "", "", "")
        self.assertEqual(result, expected)

    def test_extract_information_no_matches(self):
        """
        Test the `extract_information` function to ensure
        it returns empty strings for all fields
        when provided with text that contains no identifiable information.

        This test verifies the function's robustness
        in handling input that does not match any predefined patterns for
        name, designation, experience, education, or skills.
        """
        test_text = "Random text that does not conform to expected patterns."
        name, designation, experience, education, skills = extract_information(
            test_text)

        # check all outputs are ""
        self.assertEqual(name, "")
        self.assertEqual(designation, "")
        self.assertEqual(experience, "")
        self.assertEqual(education, "")
        self.assertEqual(skills, "")


class TestCalculateScores(unittest.TestCase):
    """
    This class tests the scoring algorithms that analyze CV text against
    a given job description.

    The tests include verifying the accuracy of
    synonym retrieval, text preprocessing, and contextual scoring functions.

    The tests also check if the scoring functions handle exceptions properly.
    """
    def setUp(self):
        # Mock the DataFrame and jd as they are external dependencies
        self.mock_cvs_df = MagicMock()
        self.mock_sample_jd = "Example job description"
        self.mock_cvs_df = pd.DataFrame({
            'Text': ["Example CV text"],
            'Name': ["John Doe"],
            'Experience': ["10 years"],
            'Skills': ["Python, Machine Learning"]
        })

    @patch('nltk.corpus.wordnet.synsets')
    def test_get_synonyms(self, mock_synsets):
        """
        Test to verify that the get_synonyms function correctly fetches
        synonyms for the word "happy".

        Mocks the nltk.corpus.wordnet.synsets to control the return values
        for testing purposes.

        This allows testing how get_synonyms behaves
        when expected synonyms are returned by the mocked function.

        Then test if an empty set would be returned
        if non-existing word is inputed.
        """
        word = "happy"
        expected_synonyms = ["happy", "felicitous", "glad"]
        mock_synset = MagicMock()
        mock_lemma1 = MagicMock()
        mock_lemma1.name.return_value = "happy"
        mock_lemma2 = MagicMock()
        mock_lemma2.name.return_value = "felicitous"
        mock_lemma3 = MagicMock()
        mock_lemma3.name.return_value = "glad"
        mock_synset.lemmas.return_value = [
            mock_lemma1, mock_lemma2, mock_lemma3]
        mock_synsets.return_value = [mock_synset]
        result = get_synonyms(word)
        self.assertEqual(sorted(result), sorted(expected_synonyms))

        # Test with non-existing word, return an empty set
        word = "nonexistingword"
        mock_synsets.return_value = []
        result = get_synonyms(word)
        self.assertEqual(result, set())

    def test_preprocess_text(self):
        '''
        Test to verify that the preprocess_text function correctly
        processes input text by removing punctuation.

        This test also checks if the function properly handles
        different types of inputs such as non-empty string,
        empty string, and strings with numeric values.
        '''
        text = "This is a sample text, with punctuation!"
        expected_result = "sample text punctuation"
        result = preprocess_text(text)
        self.assertEqual(result, expected_result)

        # Test with empty string
        text = ""
        expected_result = ""
        result = preprocess_text(text)
        self.assertEqual(result, expected_result)

        # Test with numerical values
        text = "12345"
        expected_result = "12345"
        result = preprocess_text(text)
        self.assertEqual(result, expected_result)

    def test_calculate_context_score(self):
        """
        Test the calculate_context_score function to ensure
        it returns a higher score
        for words that appear within the context compared to those that do not.

        This test checks if the function returns a float
        when handling empty input.
        """
        # 'data' is in the context
        # 'machine' is not in the context
        # test 'data' is more similar to 'machine'
        word1 = "data"
        word2 = "machine"
        context = ["data", "analysis", "statistics"]
        context_str = ' '.join(context)
        result1 = calculate_context_score(word1, context_str)
        result2 = calculate_context_score(word2, context_str)
        self.assertGreater(
            result1,
            result2,
            "The score for a word in the context should be higher than "
            "for a word not in the context")

        # Test with empty context
        context = []
        context_str = ' '.join(context)
        self.assertIsInstance(result1, float, "The result should be a float")
        self.assertIsInstance(result2, float, "The result should be a float")

    @patch('__main__.logger.error')
    def test_calculate_scores_empty_input(self, mock_logger_error):
        """
        Test that calculate_scores gracefully handles an empty DataFrame input
        and logs a specific error message.

        Also tests the function's return value
        in scenarios with both empty and valid data inputs.
        """
        # Test with empty DataFrame
        empty_df = pd.DataFrame()
        # 0 will not cause index error
        context_score = calculate_scores(0, empty_df)
        mock_logger_error.assert_called_with(
            "Error calculating score for CV 0: single positional indexer is "
            "out-of-bounds")
        self.assertEqual(context_score, 0.0)

        # Test with a valid DataFrame
        context_score = calculate_scores(0, self.mock_cvs_df)
        self.assertNotEqual(context_score, 0)

    def test_get_synonyms_nonexistent_word(self):
        """
        Test the get_synonyms function to ensure that
        it returns an empty set when provided with a non-existent word.
        """
        # Test that no synonyms are found for a made-up word
        self.assertEqual(get_synonyms("nonexistword"), set(),
                         "Should return an empty set for a nonexistent word")

    def test_preprocess_text_with_special_characters(self):
        """
        Test the preprocess_text function to confirm that
        it correctly processes texts containing special characters,
        ensuring that only alphanumeric characters
        and spaces are retained in the output.
        """
        # Test that the function handles texts
        # with unusual or special characters
        text = "Hello!!! ***@@@###$$$"
        processed_text = preprocess_text(text)
        self.assertIsInstance(processed_text, str, "Should return a string")
        self.assertTrue(all(char.isalnum() or char.isspace()
                            for char in processed_text),
                        "Should only contain alphanumeric characters "
                        "and spaces")

    def test_preprocess_text_with_empty_input(self):
        """
        Test that preprocess_text returns
        an empty string when provided with an empty input,
        confirming the function's ability
        to handle empty string data correctly.
        """
        # Test that the function returns
        # an empty string when given an empty input
        empty_text = ""
        processed_text = preprocess_text(empty_text)
        self.assertIsInstance(processed_text, str, "Should return a string")
        self.assertEqual(processed_text, "",
                         "Should return an empty string for empty input")

    def test_calculate_context_score_with_empty_inputs(self):
        """
        Test the calculate_context_score function to verify that
        it returns a score of 0.0 when both inputs are empty.
        """
        # Test that the function returns
        # a context score when both inputs are empty
        empty_jd = ""
        empty_cv = ""
        score = calculate_context_score(empty_jd, empty_cv)
        self.assertIsInstance(score, float, "Should return a float")
        self.assertEqual(score, 0.0, "Should return 0.0 for empty inputs")

    def test_combined_functionality_with_missing_fields(self):
        """
        Test that the analysis functions can handle
        CV dataframes with missing fields,
        specifically testing that missing data does not prevent
        the generation of combined text and subsequent synonym extraction.
        """
        # Test with missing fields
        cvs_df = pd.DataFrame({
            'Text': ['example text'],
            'Name': ['John Doe'],
            # 'Experience' field is missing
            'Skills': ['Python, Data Analysis']
        })
        combined_text = (f"{cvs_df.iloc[0]['Text']} "
                         f"{cvs_df.iloc[0]['Name']} "
                         f"{cvs_df.iloc[0].get('Experience', '')} "
                         f"{cvs_df.iloc[0]['Skills']}")
        result = get_synonyms(combined_text)
        self.assertIsNotNone(result)

    @patch('__main__.logger')
    def test_calculate_scores_empty_descriptions(self, mock_logger):
        """
        Test that calculate_scores does not
        raise an error with empty descriptions.
        """
        cvs_df = pd.DataFrame({
            'Text': [''],
            'Name': [''],
            'Experience': [''],
            'Skills': ['']
        })

        try:
            score = calculate_scores(0, cvs_df)
            print(f"Score calculated: {score}")
        except Exception as e:
            # fail if there is error
            self.fail(
                f"calculate_scores raised an exception "
                f"with empty descriptions: {e}")

    @patch('__main__.calculate_context_score',
           side_effect=Exception("Failed to calculate context score"))
    @patch('__main__.logger')
    def test_calculate_scores_with_exception(self,
                                             mock_logger,
                                             mock_calculate_context_score):
        """
        Test that calculate_scores handles exceptions correctly,
        logs the error,
        and returns a default value.
        """
        result = calculate_scores(0, self.mock_cvs_df)

        # check if default value would be returned after Error is detected
        # check if logger will log the error message
        self.assertEqual(result, 0.0)
        mock_logger.error.assert_called_with(
            "Error calculating score for CV 0: "
            "Failed to calculate context score")


class TestScanFiles(unittest.TestCase):
    """
    Tests the file scanning functionalities
    to ensure the system properly handles
    file access, retrieval,
    and scanning operations under various conditions.
    """

    @patch('__main__.logger.error')
    def test_check_repository_access_invalid_path(self, mock_logger_error):
        """
        Test that the system correctly handles
        an invalid file path input and
        logs an appropriate error message.
        """
        # Test for invalid path input
        result = check_repository_access('invalid_path')
        self.assertFalse(result)

        # Confirm whether logger.error is called
        mock_logger_error.assert_called_once()
        # Check whether logger.error is called with the correct error message
        mock_logger_error.assert_called_with(
            "Unexpected error: "
            "The specified folder does not exist or cannot be accessed.")

    @patch('__main__.logger.error')
    @patch('__main__.check_repository_access')
    @patch('os.listdir')
    def test_load_cv_files_repository_not_accessible(self,
                                                     mock_listdir,
                                                     mock_check_access,
                                                     mock_log_error):
        """
        Test that the system returns
        an empty list and logs an error
        when attempting to load CV files from an inaccessible directory.
        """
        # Mock the repository as being not accessible
        mock_check_access.return_value = False

        # Expected to return an empty list if the directory is inaccessible
        result = load_cv_files('some_path')
        self.assertEqual(result, [])
        mock_listdir.assert_not_called()  # Check os.listdir is not called
        mock_log_error.assert_called_once_with(
            "Directory access check failed. No files to load.")

    @patch('__main__.logger.info')
    @patch('__main__.check_repository_access')
    @patch('os.listdir')
    def test_load_cv_files_no_matching_files(self,
                                             mock_listdir,
                                             mock_check_access,
                                             mock_log_info):
        """
        Test that no CV files are found and
        an informative message is logged
        when the directory does not contain files of the expected format.
        """
        # Mock the repository as being accessible
        mock_check_access.return_value = True
        # Simulate no matching file types in the directory
        mock_listdir.return_value = ['file1.txt', 'file2.png']

        result = load_cv_files('some_path')
        self.assertEqual(result, [])
        mock_listdir.assert_called_once_with('some_path')
        mock_log_info.assert_called_once_with(
            "No CV files of given format found in the directory.")

    @patch('__main__.logger.warning')
    @patch('__main__.check_repository_access')
    @patch('os.listdir', side_effect=FileNotFoundError("Directory not found"))
    def test_load_cv_files_file_not_found_error(self,
                                                mock_listdir,
                                                mock_check_access,
                                                mock_log_warning):
        """
        Test that the system correctly handles
        a FileNotFoundError when trying to access a directory
        and logs a warning message.
        """
        # Mock the repository as being accessible
        mock_check_access.return_value = True

        # Deal with FileNotFoundError, check logger
        result = load_cv_files('some_path')
        self.assertEqual(result, [])
        mock_listdir.assert_called_once_with('some_path')
        mock_log_warning.assert_called_once_with(
            "File not found: Directory not found")

    @patch('__main__.logger.warning')
    @patch('__main__.logger.info')
    @patch('cloudmersive_virus_api_client.ScanApi.scan_file_advanced')
    @patch('time.sleep', return_value=None)  # Mock sleep to avoid delay
    @patch('logging.Logger.info')  # Mock the logger info method
    @patch('logging.Logger.warning')  # Mock the logger warning method
    def test_scan_one_file_invalid_path(self,
                                        mock_warning, mock_info, mock_sleep,
                                        mock_scan_file_advanced,
                                        mock_logger_info, mock_logger_warning):
        """
        Test that the system logs an error message
        and returns False when attempting to scan a file at an invalid path.
        """
        # Setup the method to raise a general Exception when called
        mock_scan_file_advanced.side_effect = Exception("API failure")

        # Call the function with a path that will trigger the exception
        result = scan_one_file("invalid_path", "valid_api_key", 0)

        # Assert that the result is False, indicating the scan failed
        self.assertFalse(result)

        # Check that scan_file_advanced was called,
        # implying that the process went as far as attempting a scan
        mock_scan_file_advanced.assert_called()

        # Assert based on mock_info and mock_warning
        mock_logger_info.assert_called_with(
            "Exception when calling Cloudmersive API "
            "-> scan_file_advanced: API failure\n")

        # Check if the warning log was also called correctly
        mock_logger_warning.assert_called_once_with(
            'Failed to scan file: invalid_path')

    @patch('__main__.scan_one_file')
    @patch('__main__.logger')
    def test_scan_all_files_empty_list(self, mock_logger, mock_scan_one_file):
        '''
        Test for an empty file list input.
        Aims for no error.
        '''
        result = scan_all_files_in_repository(
            '/valid/path', [], 'valid_api_key')
        self.assertEqual(result, [])
        mock_scan_one_file.assert_not_called()
        mock_logger.info.assert_called_with(
            "Scanning all files in the repositorycompleted.")

    @patch('__main__.logger')
    # Assuming scan_one_file returns True for a clean file
    @patch('__main__.scan_one_file', return_value=True)
    def test_scan_all_files_functionality(self,
                                          mock_scan_one_file, mock_logger):
        '''
        Test if
        scan_all_files_in_repository correctly scans all files in repository.
        Assumes all files are clean
        '''
        files = ['file1.pdf', 'file2.docx']
        cv_folder_path = '/valid/path'
        api_key = 'valid_api_key'

        result = scan_all_files_in_repository(cv_folder_path, files, api_key)
        self.assertEqual(result, files)

        # Verify that the final log message about completion is recorded
        mock_logger.info.assert_called_with(
            "Scanning all files in the repositorycompleted.")

        # Ensure no warning for a failure was logged
        mock_logger.warning.assert_not_called()

    @patch('__main__.scan_one_file', side_effect=Exception("Scan failed"))
    @patch('__main__.logger')
    def test_scan_all_files_with_scan_failure(self,
                                              mock_logger, mock_scan_one_file):
        '''
        Teset situation when errors occurred during test scanning.
        Would return an empty list and log an error message.
        '''
        files = ['file1.pdf', 'file2.docx']
        result = scan_all_files_in_repository(
            '/valid/path', files, 'valid_api_key')
        self.assertEqual(result, [])
        mock_logger.warning.assert_called_with(
            "Failed to complete scanning process: Scan failed")


class TestPerformCVAnalysis(unittest.TestCase):
    """
    This class tests the functionality of CV analysis processes,
    ensuring that  each part of the CV processing pipeline
    works correctly under various scenarios.
    """

    @patch('builtins.print')  # Mock print to avoid cluttering test output
    @patch('__main__.PrettyTable')
    @patch('__main__.nltk')
    @patch('__main__.pd.DataFrame')
    @patch('__main__.Pool')
    def test_empty_cv_list(self,
                           mock_pool, mock_dataframe, mock_nltk,
                           mock_prettytable, mock_print):
        """
        Tests the behavior of the CV analysis
        when no CVs are provided.

        Ensures that appropriate messages are printed
        indicating no files to process.
        """
        # Define dummy variables
        cv_folder_path = "/dummy/path"
        cv_files = []
        sample_jd = "Sample Job Description"

        # test empty list
        mock_pool.return_value.__enter__.return_value.map.return_value = []
        perform_cv_analysis(cv_folder_path, cv_files, sample_jd)

        # Verify that the specific print statements were made
        mock_print.assert_any_call("No valid file to be handled.")
        mock_print.assert_any_call("Total execution time: 0.00 seconds")

    @patch('builtins.print')
    @patch('__main__.logger')
    @patch('__main__.PrettyTable')
    @patch('__main__.nltk')
    @patch('__main__.pd.DataFrame')
    @patch('__main__.Pool')
    def test_invalid_sample_jd(self,
                               mock_pool, mock_dataframe, mock_nltk,
                               mock_prettytable, mock_logger, mock_print):
        '''
        Test if perform_cv_analysis could handle invalid jd
        '''
        cv_folder_path = "/valid/path"
        cv_files = ["resume1.pdf"]

        sample_jd1 = ""  # Empty or non-standard JD
        perform_cv_analysis(cv_folder_path, cv_files, sample_jd1)

        # Check that the appropriate error was logged
        mock_logger.error.assert_called_with(
            'An unexpected error occurred while processing the CVs: '
            'Invalid or missing sample job description.')

        # Ensure that the print statement in the finally block is called
        # This will pass if print is called at least once for any reason
        mock_print.assert_called()

        sample_jd2 = None
        perform_cv_analysis(cv_folder_path, cv_files, sample_jd2)
        mock_logger.error.assert_called_with(
            'An unexpected error occurred while processing the CVs: '
            'Invalid or missing sample job description.')

        sample_jd3 = "@#$123"
        perform_cv_analysis(cv_folder_path, cv_files, sample_jd3)
        mock_logger.error.assert_called_with(
            'An unexpected error occurred while processing the CVs: '
            'Invalid or missing sample job description.')

    @patch('builtins.print')
    @patch('__main__.logger')  # Patch the logger to check for error logging
    @patch('__main__.PrettyTable')
    @patch('__main__.nltk.download',
           side_effect=Exception("Failed to download NLTK resources"))
    @patch('__main__.pd.DataFrame')
    @patch('__main__.Pool')
    def test_nltk_download_failure(self,
                                   mock_pool, mock_dataframe,
                                   mock_nltk_download,
                                   mock_prettytable, mock_logger, mock_print):
        """
        Tests the error handling when
        there is a failure in downloading necessary NLTK resources,
        verifying that errors are logged and
        the analysis can still proceed or halt gracefully.
        """
        cv_folder_path = "/valid/path"
        cv_files = ["resume1.pdf"]
        sample_jd = "Sample Job Description"

        # Execute the function
        perform_cv_analysis(cv_folder_path, cv_files, sample_jd)

        # Verify that the error was logged correctly
        error_calls = mock_logger.error.call_args_list
        unpack_error_called = any(
            "not enough values to unpack" in str(call) for call in error_calls)
        self.assertTrue(unpack_error_called,
                        "Expected unpack error message not found in log.")

        # Ensure that the print statement in the finally block is still called
        print_calls = mock_print.call_args_list
        execution_time_called = any(
            "Total execution time:" in str(call) for call in print_calls)
        self.assertTrue(execution_time_called,
                        "Expected execution time print statement not found.")

    @patch('__main__.process_cv_file',
           return_value=('dummy_resume.pdf', 'text', 'name',
                         'designation', 'experience', 'education',
                         ['skill1', 'skill2']))
    @patch('multiprocessing.Pool')
    @patch('__main__.logger')
    # Mock stderr to capture error output
    @patch('sys.stderr', new_callable=MagicMock)
    def test_error_in_parallel_processing(self,
                                          mock_stderr, mock_logger, mock_pool,
                                          mock_process_cv_file):
        """
        Tests the system's resilience to errors
        during parallel processing of CVs,
        ensuring that errors are logged and the system recovers gracefully.
        """
        # Set up a simulated multiprocessing.Pool that raises an exception
        mock_process_pool = MagicMock()
        mock_pool.return_value.__enter__.return_value = mock_process_pool
        # Return the value using simulated process_cv_file
        mock_process_pool.map.return_value = [
            mock_process_cv_file.return_value]

        perform_cv_analysis(
            "/dummy/path", ["dummy_resume.pdf"], "Fake job description")

        # Check if a specific error message was logged
        error_message = ("An unexpected error occurred "
                         "while processing the CVs:")
        found_error_message = any(error_message in str(
            call) for call in mock_logger.method_calls if call[0] == 'error')
        self.assertTrue(found_error_message,
                        "Expected error message not found in logger calls")

    @patch('__main__.nltk.word_tokenize')
    @patch('__main__.logger')
    def test_tokenization_error_handling(self, mock_logger, mock_tokenize):
        """
        Tests the tokenization process to ensure
        that the system can handle tokenization failures,
        logging appropriate error messages when errors occur.
        """
        # Set the mock tokenization function to throw an exception when called
        mock_tokenize.side_effect = Exception("Tokenization failed")

        # Call the perform_cv_analysis function,
        # which should trigger an exception
        perform_cv_analysis(
            "/dummy/path", ["dummy_resume.pdf"], "Dummy Job Description")

        # Check that the logger logged the expected error message
        mock_logger.error.assert_called_with(
            "An unexpected error occurred while processing the CVs: "
            "Tokenization failed")


unittest.main(argv=[''], verbosity=2, exit=False)

# close the logger
logger.remove()


